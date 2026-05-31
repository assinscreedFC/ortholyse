# =============================================================================
# Auteur  : HAMMOUCHE Anis
# Email   : anis.hammouche@etu.u-paris.fr
# Version : 1.0
# =============================================================================
import json
import csv
from datetime import datetime
import os
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from fpdf import FPDF #type:ignore
from docx import Document #type:ignore

def exporte_json(chemin: str, data: dict) -> None:
    """
    Exporte les données en format JSON.
    """
    with open(chemin, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)
    #print(f"Données exportées en JSON : {chemin}")

def exporte_txt(chemin: str, data: dict) -> None:
    """
    Exporte les données en format texte (.txt).
    Chaque clé-valeur sera écrite ligne par ligne.
    """
    with open(chemin, "w", encoding="utf-8") as f:
        for cle, valeur in data.items():
            f.write(f"{cle}: {valeur}\n")
    #print(f"Données exportées en TXT : {chemin}")

def exporte_csv_column(chemin: str, data: dict) -> None:
    """
    Exporte les données en format CSV.
    Les clés du dictionnaire seront utilisées comme en-tête.
    """
    with open(chemin, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Clé", "Valeur"])  # En-tête
        for cle, valeur in data.items():
            writer.writerow([cle, valeur])  # Chaque ligne correspond à clé-valeur
    #print(f"Données exportées en CSV : {chemin}")

def exporte_csv_row(chemin: str, data: dict) -> None:
    """
    Exporte les données en format CSV sous forme de tableau.
    Les clés du dictionnaire sont utilisées comme en-têtes des colonnes.
    """
    with open(chemin, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        
        # Écrit l'en-tête : les clés du dictionnaire
        writer.writerow(data.keys())
        
        values = [v if isinstance(v, list) else [v] for v in data.values()] #pour transformer tout les element en liste

        # Transposer les colonnes en lignes
        rows = zip(*values) #l'etoile sert a decompacter les element cette a dire [[1,2,3],[a,b,c]] devient [1,2,3],[a,b,c]
        writer.writerows(rows)  # Écrit les lignes



def exporte_pdf(chemin, data, titre, font_dir='../app/assets/Fonts/Poppins'):
    """
    Exporte les données en format PDF avec :
    - un titre centré
    - un tableau des données
    - une section séparée pour le texte de la transcription (si présent)

    :param chemin: Chemin de sortie du fichier PDF
    :param data: Dictionnaire des données à exporter
    :param titre: Titre du document
    :param font_dir: Répertoire contenant les polices (.ttf)
    """
    pdf = FPDF()
    pdf.add_page()

    # Chargement des polices
    font_dir = os.path.abspath(font_dir)
    pdf.add_font('Poppins-Bold', '', os.path.join(font_dir, 'Poppins-Bold.ttf'), uni=True)
    pdf.add_font('Poppins', '', os.path.join(font_dir, 'Poppins-Regular.ttf'), uni=True)

    # Titre principal
    pdf.set_font('Poppins-Bold', '', 24)
    pdf.cell(0, 12, txt=titre, ln=True, align='C')
    pdf.ln(10)

    # Date et heure
    now = datetime.now()
    date_str = now.strftime("%d/%m/%Y")
    time_str = now.strftime("%H:%M")

    # Configuration tableau
    pdf.set_line_width(0.8)
    padding_horizontal = 20
    table_width = pdf.w - 2 * padding_horizontal
    col_width = table_width / 2
    row_height = 14
    start_x = padding_horizontal

    def table_row(key, value):
        pdf.set_x(start_x)
        pdf.set_font('Poppins-Bold', '', 14)
        pdf.cell(col_width, row_height, str(key), border=1, align='C')
        pdf.set_font('Poppins', '', 14)
        pdf.cell(col_width, row_height, str(value), border=1, ln=1, align='C')

    # Données du tableau
    table_row("Date", date_str)
    table_row("Heure", time_str)
    for cle, valeur in data.items():
        if cle != "texte":
            table_row(cle, valeur)

    # Ajout de la transcription si elle existe
    if "texte" in data:
        pdf.ln(15)
        # Titre H1 pour la transcription
        pdf.set_font('Poppins-Bold', '', 20)
        pdf.cell(0, 12, txt="Texte de la transcription", ln=True, align='L')

        pdf.ln(5)
        # Paragraphe
        pdf.set_font('Poppins', '', 12)
        pdf.multi_cell(0, 8, txt=data["texte"])


    # Sauvegarde
    pdf.output(chemin)




def exporte_docx(chemin, data, titre, font_path='../app/assets/Fonts/Poppins-Regular.ttf'):
    doc = Document()

    # Style global
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Poppins'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:ascii'), 'Poppins')
    style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Poppins')

    # Titre principal
    titre_paragraph = doc.add_paragraph()
    titre_run = titre_paragraph.add_run(titre)
    titre_run.bold = True
    titre_run.font.size = Pt(24)
    titre_run.font.name = 'Poppins'
    titre_run.font.color.rgb = RGBColor(0, 0, 0)
    titre_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Espace

    now = datetime.now()
    date_str = now.strftime("%d/%m/%Y")
    time_str = now.strftime("%H:%M")

    # Création du tableau
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def set_cell_height(cell, height_twips):
        tr = cell._tc.getparent()
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(height_twips))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    def add_row(cle, valeur):
        row_cells = table.add_row().cells

        # Cellule gauche (clé, en gras)
        cell_key = row_cells[0]
        para_key = cell_key.paragraphs[0]
        run_key = para_key.add_run(str(cle))
        run_key.font.name = 'Poppins'
        run_key.font.size = Pt(14)
        run_key.font.color.rgb = RGBColor(0, 0, 0)
        run_key.bold = True
        para_key.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_key.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_height(cell_key, 600)

        # Cellule droite (valeur)
        cell_val = row_cells[1]
        para_val = cell_val.paragraphs[0]
        run_val = para_val.add_run(str(valeur))
        run_val.font.name = 'Poppins'
        run_val.font.size = Pt(14)
        run_val.font.color.rgb = RGBColor(0, 0, 0)
        para_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_height(cell_val, 600)

    # Contenu du tableau
    add_row("Date", date_str)
    add_row("Heure", time_str)

    texte_transcription = None
    for cle, valeur in data.items():
        if cle == "texte":
            texte_transcription = valeur
        else:
            add_row(cle, valeur)

    # Ajout du texte de transcription (en dehors du tableau)
    if texte_transcription:
        doc.add_paragraph()  # Espace

        # Titre H1
        transcription_title = doc.add_paragraph()
        run_title = transcription_title.add_run("Texte de la transcription")
        run_title.font.name = 'Poppins'
        run_title.font.size = Pt(20)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(0, 0, 0)

        # Texte en paragraphe
        para = doc.add_paragraph()
        run = para.add_run(texte_transcription)
        run.font.name = 'Poppins'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Sauvegarde

    doc.save(chemin)

data={
    "nom": "Alice",
    "âge": 25,
    "profession": "Développeuse",
    "ville": "Paris",
    "texte":"anis a 25ans il est a paris et est trés trés riche mais il sais toujours pas ecririr"
}
#exporte_docx('./',data,"analyse°1")
